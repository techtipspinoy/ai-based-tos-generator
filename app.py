import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

# --- Helper: Bloom's Taxonomy Distribution (default for 30 items) ---
BLOOM_DISTRIBUTION = {
    "Remembering": 0.20,
    "Understanding": 0.20,
    "Applying": 0.25,
    "Analyzing": 0.15,
    "Evaluating": 0.10,
    "Creating": 0.10
}

# --- Bloom's Verbs for Item Writing ---
BLOOM_VERBS = {
    "Remembering": ["define", "list", "name", "recall", "identify"],
    "Understanding": ["explain", "describe", "summarize", "paraphrase", "classify"],
    "Applying": ["solve", "use", "apply", "demonstrate", "calculate"],
    "Analyzing": ["compare", "contrast", "categorize", "differentiate", "infer"],
    "Evaluating": ["justify", "assess", "critique", "defend", "recommend"],
    "Creating": ["design", "construct", "develop", "propose", "formulate"]
}

# --- Fake but realistic sample MELC bank (for demo & validation) ---
MELC_DATABASE = {
    "Mathematics": {
        "Grade 7": {
            "Q1": [
                {"code": "M7NS-Ia-1", "desc": "Describes well-defined sets, subsets, universal set, and the null set and cardinality of sets."},
                {"code": "M7NS-Ib-1", "desc": "Solves problems involving sets."},
                {"code": "M7NS-Ic-1", "desc": "Represents the union and intersection of sets using Venn Diagrams."},
            ],
            "Q2": [
                {"code": "M7NS-IIa-1", "desc": "Expresses rational numbers from fraction form to decimal form and vice versa."},
                {"code": "M7NS-IIc-1", "desc": "Performs operations on rational numbers."},
            ]
        },
        "Grade 8": {
            "Q1": [
                {"code": "M8AL-Ia-b-1", "desc": "Factors polynomials."},
                {"code": "M8AL-Ic-1", "desc": "Illustrates rational algebraic expressions."},
            ]
        }
    },
    "Science": {
        "Grade 8": {
            "Q3": [
                {"code": "S8FE-IIIa-15", "desc": "Explain how different types of waves form and behave."},
                {"code": "S8FE-IIIb-16", "desc": "Describe how waves carry energy."},
                {"code": "S8FE-IIIc-17", "desc": "Relate the characteristics of waves to their behavior (e.g., reflection, refraction)."},
            ]
        }
    },
    "English": {
        "Grade 9": {
            "Q3": [
                {"code": "EN9RC-IIIa-13.10", "desc": "Evaluate information to expand, review, or update knowledge."},
                {"code": "EN9RC-IIIb-14.3", "desc": "React critically to the ideas presented in a text."},
            ]
        }
    }
}

# --- Utility: Add heading with style ---
def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- Utility: Set cell borders ---
def set_cell_borders(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ["top", "left", "bottom", "right"]:
        value = kwargs.get(border_name, "single")
        if value:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), value)
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tcBorders.append(border)
    tcPr.append(tcBorders)

# --- Generate TOS DataFrame ---
def generate_tos(competencies, total_items=30):
    # Compute item count per cognitive level
    counts = {level: round(BLOOM_DISTRIBUTION[level] * total_items) for level in BLOOM_DISTRIBUTION}
    
    # Adjust to ensure sum = total_items (rounding may cause ±1 error)
    diff = total_items - sum(counts.values())
    if diff > 0:
        counts["Applying"] += diff
    elif diff < 0:
        counts["Creating"] += diff  # reduce least-weighted

    rows = []
    item_no = 1
    
    # Distribute items across competencies & Bloom's levels
    for comp in competencies:
        # Assign ~equal share (or more to earlier ones if uneven)
        comp_items = total_items // len(competencies)
        extra = 1 if competencies.index(comp) < total_items % len(competencies) else 0
        comp_items += extra
        
        # Distribute this comp’s items across Bloom's levels
        comp_counts = {level: round(counts[level] * comp_items / total_items) for level in counts}
        comp_diff = comp_items - sum(comp_counts.values())
        if comp_diff > 0:
            comp_counts["Applying"] += comp_diff
        
        for level, n in comp_counts.items():
            for _ in range(n):
                # Determine item type by Bloom's level
                if level in ["Remembering", "Understanding"]:
                    item_type = "Multiple Choice"
                elif level in ["Applying", "Analyzing"]:
                    item_type = "Short Answer"
                else:  # Evaluating, Creating
                    item_type = "Essay"
                
                point_val = 1 if item_type == "Multiple Choice" else (2 if item_type == "Short Answer" else 5)
                
                rows.append({
                    "Item No.": item_no,
                    "Cognitive Level": level,
                    "Competency (MELC)": f"{comp['code']}: {comp['desc']}",
                    "Item Type": item_type,
                    "Point Value": point_val,
                    "Remarks": ""
                })
                item_no += 1
    return pd.DataFrame(rows)

# --- Generate quiz items ---
def generate_quiz_items(tos_df):
    items = []
    for _, row in tos_df.iterrows():
        comp_code = row["Competency (MELC)"].split(":")[0].strip()
        level = row["Cognitive Level"]
        item_type = row["Item Type"]
        points = row["Point Value"]
        
        verb = BLOOM_VERBS[level][0].capitalize()
        topic = "waves" if "S8FE" in comp_code else "algebraic expressions" if "M8AL" in comp_code else "information"
        
        item_text = ""
        answer = ""
        
        if item_type == "Multiple Choice":
            item_text = f"{verb} the following:\nWhat is the primary characteristic of mechanical waves?\n"
            item_text += "A. They can travel through vacuum.\n"
            item_text += "B. They require a medium to propagate.\n"
            item_text += "C. They are always transverse.\n"
            item_text += "D. They travel faster than light.\n"
            answer = "✓ B"
        elif item_type == "Short Answer":
            item_text = f"{verb} how sound waves carry energy through air.\n"
            item_text += "(2–3 sentences)\n"
            answer = "[Sample] Sound waves carry energy by compressing and rarefying air particles, transferring kinetic energy from one particle to the next."
        else:  # Essay
            item_text = f"{verb} an experiment to demonstrate wave reflection and refraction using everyday materials. Justify your design choices.\n"
            answer = (
                "[Rubric: 5 pts total]\n"
                "• Content (3 pts): Accurate science, clear steps\n"
                "• Organization (1 pt): Logical flow\n"
                "• Mechanics (1 pt): Grammar, spelling\n"
            )
        
        items.append({
            "item_no": row["Item No."],
            "text": item_text,
            "answer": answer,
            "points": points
        })
    return items

# --- Create Word Document ---
def create_word_doc(tos_df, quiz_items, metadata):
    doc = Document()
    
    # === Cover Page ===
    section = doc.sections[0]
    section.page_height = Inches(13)
    section.page_width = Inches(8.5)
    section.orientation = WD_ORIENT.PORTRAIT
    
    add_heading(doc, "DEPARTMENT OF EDUCATION", 1)
    add_heading(doc, "TABLE OF SPECIFICATIONS & ASSESSMENT TOOL", 1)
    doc.add_paragraph("\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"School: {metadata['school']}\n").bold = True
    p.add_run(f"Grade Level: {metadata['grade']}\n")
    p.add_run(f"Subject: {metadata['subject']}\n")
    p.add_run(f"Quarter: {metadata['quarter']}\n")
    p.add_run(f"Prepared by: {metadata['teacher']}\n")
    p.add_run(f"Date: {metadata['date']}\n")
    doc.add_page_break()
    
    # === TOS Table ===
    add_heading(doc, "I. TABLE OF SPECIFICATIONS", 2)
    
    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ["Item No.", "Cognitive Level", "Competency (MELC)", "Item Type", "Point Value", "Remarks"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_borders(hdr_cells[i])
    
    # Populate rows
    for _, row in tos_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Item No."])
        cells[1].text = row["Cognitive Level"]
        cells[2].text = row["Competency (MELC)"]
        cells[3].text = row["Item Type"]
        cells[4].text = str(row["Point Value"])
        cells[5].text = row["Remarks"]
        for cell in cells:
            set_cell_borders(cell)
    
    doc.add_paragraph("\n")
    
    # === Quiz ===
    add_heading(doc, "II. QUIZ / EXAMINATION", 2)
    doc.add_paragraph(f"General Instructions: Answer the following. Total Points: {tos_df['Point Value'].sum()}", style='Intense Quote')
    
    for item in quiz_items:
        p = doc.add_paragraph()
        p.add_run(f"{item['item_no']}. ").bold = True
        p.add_run(item["text"])
    
    # === Answer Key ===
    doc.add_page_break()
    add_heading(doc, "III. ANSWER KEY & RUBRICS", 2)
    for item in quiz_items:
        p = doc.add_paragraph()
        p.add_run(f"{item['item_no']}. ").bold = True
        p.add_run(item["answer"])
    
    # Footer
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = "Aligned with the Most Essential Learning Competencies (MELCs) – DepEd Order No. 012, s. 2023"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

# ======================
# STREAMLIT UI
# ======================
st.set_page_config(
    page_title="DepEd TOS & Quiz Generator",
    page_icon="",
    layout="wide"
)

st.title("DepEd Table of Specifications & Quiz Generator")
st.markdown(
    "> *Aligned with K to 12 MELCs (2023–2024) | For Grades 7–10*"
)

# Sidebar for metadata
with st.sidebar:
    st.header("School Information")
    school = st.text_input("School Name", "Tiring  National High School")
    teacher = st.text_input("Prepared by (Teacher)", "Juan Dela Cruz")
    date = st.date_input("Date", format="YYYY-MM-DD")

# Main form
col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("📋 Input Assessment Details")
    
    grade = st.selectbox("Grade Level", ["Grade 7", "Grade 8", "Grade 9", "Grade 10"])
    subject = st.selectbox("Subject", list(MELC_DATABASE.keys()))
    
    # Dynamic quarter options
    quarters = list(MELC_DATABASE.get(subject, {}).get(grade, {}).keys())
    quarter = st.selectbox("Quarter", quarters if quarters else ["Q1", "Q2", "Q3", "Q4"])
    
    # Competency selection (multi-select from DB or free text)
    st.markdown("#### Select or Enter Competencies")
    use_sample = st.checkbox("✅ Use sample MELCs (recommended)", value=True)
    
    competencies = []
    if use_sample and subject in MELC_DATABASE and grade in MELC_DATABASE[subject] and quarter in MELC_DATABASE[subject][grade]:
        options = [
            f"{c['code']}: {c['desc']}" 
            for c in MELC_DATABASE[subject][grade][quarter]
        ]
        selected = st.multiselect(
            "Choose MELCs",
            options,
            default=options[:2] if options else []
        )
        for s in selected:
            code, desc = s.split(": ", 1)
            competencies.append({"code": code, "desc": desc})
    else:
        st.info("Enter competencies manually (one per line, format: `CODE: Description`)")
        raw = st.text_area("Competencies (MELC format)", height=120)
        if raw.strip():
            for line in raw.strip().split("\n"):
                if ":" in line:
                    code, desc = line.split(":", 1)
                    competencies.append({"code": code.strip(), "desc": desc.strip()})
    
    total_items = st.slider("Total Number of Items", min_value=10, max_value=50, value=30, step=5)

with col2:
    st.subheader("⚙️ Preview & Generate")
    st.write(f"**Selected Competencies:** {len(competencies)}")
    for c in competencies[:3]:
        st.caption(f"• `{c['code']}`")
    if len(competencies) > 3:
        st.caption(f"... and {len(competencies)-3} more")
    
    st.markdown("---")
    if st.button("🚀 Generate TOS & Quiz", type="primary", use_container_width=True):
        if not competencies:
            st.error("⚠️ Please input at least one competency.")
        else:
            with st.spinner("Generating TOS and quiz..."):
                # Generate TOS
                tos_df = generate_tos(competencies, total_items)
                
                # Generate quiz items
                quiz_items = generate_quiz_items(tos_df)
                
                # Store in session state
                st.session_state.tos_df = tos_df
                st.session_state.quiz_items = quiz_items
                st.session_state.metadata = {
                    "school": school,
                    "grade": grade,
                    "subject": subject,
                    "quarter": quarter,
                    "teacher": teacher,
                    "date": str(date)
                }

# Display outputs
if "tos_df" in st.session_state:
    st.success("✅ TOS & Quiz generated!")
    
    tab1, tab2, tab3 = st.tabs(["📊 TOS (Preview)", "📝 Quiz (Preview)", "📥 Download"])
    
    with tab1:
        st.dataframe(st.session_state.tos_df, use_container_width=True)
    
    with tab2:
        for item in st.session_state.quiz_items:
            st.markdown(f"**{item['item_no']}.** {item['text']}")
            with st.expander("💡 Sample Answer/Rubric"):
                st.text(item["answer"])
    
    with tab3:
        # Generate Word doc
        doc = create_word_doc(
            st.session_state.tos_df,
            st.session_state.quiz_items,
            st.session_state.metadata
        )
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"TOS_Quiz_{st.session_state.metadata['grade']}_{st.session_state.metadata['subject']}_Q{st.session_state.metadata['quarter']}.docx"
        
        st.download_button(
            label="📥 Download Word Document (.docx)",
            data=buffer,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
        st.info("✅ Document includes: Cover page, TOS table, Quiz items, Answer key, and DepEd compliance footer.")

# Footer
st.markdown("---")
st.caption(
    "💡 Tip: For best results, use official MELC codes from [DepEd MELC Portal](https://commons.deped.gov.ph/melc). "
    "This tool does not store your data — all processing happens in your browser/session."
)